# app/core/ingest.py
from __future__ import annotations

import mimetypes
import os
import re
from typing import Dict, List, Tuple

# PDF (no pdfplumber; use pypdf)
# pip install pypdf
from pypdf import PdfReader

# DOCX
# pip install python-docx
from docx import Document


def normalize_text(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = s.replace("\x00", " ")
    s = re.sub(r"[ \t]+\n", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def extract_text_from_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text_from_pdf(path: str) -> str:
    out: List[str] = []
    reader = PdfReader(path)
    for page in reader.pages:
        txt = page.extract_text() or ""
        if txt.strip():
            out.append(txt)
    return "\n\n".join(out)


def extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    parts = []
    # paragraphs
    for p in doc.paragraphs:
        t = p.text or ""
        if t.strip():
            parts.append(t)
    # simple tables
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [normalize_text(c.text or "") for c in row.cells]
            row_text = " | ".join([c for c in cells if c.strip()])
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_text_from_any(path: str, source: str = "") -> Dict:
    """
    Returns dict: {"text": str, "meta": {...}}
    """
    mime, _ = mimetypes.guess_type(source or path)
    ext = (os.path.splitext(source or path)[1] or "").lower()

    text = ""
    if ext in [".txt", ".md", ".csv", ".log"] or (mime and mime.startswith("text/")):
        text = extract_text_from_txt(path)
    elif ext == ".pdf" or (mime == "application/pdf"):
        text = extract_text_from_pdf(path)
    elif ext in [".docx"] or (mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
        text = extract_text_from_docx(path)
    else:
        # last-ditch: try text
        try:
            text = extract_text_from_txt(path)
        except Exception:
            raise ValueError(f"Unsupported file type: {ext or mime or '<unknown>'}")

    text = normalize_text(text)
    return {
        "text": text,
        "meta": {
            "source": os.path.basename(source or path),
            "mime": mime or "text/plain",
        },
    }


def _split_sentences(text: str) -> List[str]:
    # Simple sentence splitter; avoids heavyweight deps
    # Split on period/question/exclamation + space/newline
    parts = re.split(r"(?<=[\.\?\!])\s+", text)
    return [p.strip() for p in parts if p and p.strip()]


def smart_chunk(
    text: str,
    *,
    target_tokens: int = 300,
    overlap_tokens: int = 40,
) -> List[Dict]:
    """
    Chunk by rough word-count, respecting sentence boundaries.
    Returns list of dicts: [{"text": "...", "meta": {...}}, ...]
    """
    if not text.strip():
        return []

    # rough token estimate: 1 token ~= 0.75 words for English; we just use words here
    words_per_chunk = max(40, int(target_tokens))  # keep reasonable minimum
    overlap = max(0, int(overlap_tokens))

    sentences = _split_sentences(text)
    chunks: List[str] = []
    cur: List[str] = []
    cur_len = 0

    def flush():
        nonlocal cur, cur_len
        if cur:
            chunks.append(" ".join(cur).strip())
            cur = []
            cur_len = 0

    for s in sentences:
        w = s.split()
        if cur_len + len(w) > words_per_chunk:
            # flush current chunk, then start with overlap from end
            flush()
            if overlap > 0 and chunks:
                # seed with overlap: take tail words from previous chunk
                tail = chunks[-1].split()
                cur = tail[-overlap:]
                cur_len = len(cur)
        cur.extend(w)
        cur_len += len(w)

    flush()

    return [{"text": c, "meta": {}} for c in chunks]


def ingest_file(
    user_id: str,
    path: str,
    filename: str,
    *,
    upserter,  # callable: (user_id, source, chunks, metadata) -> int
) -> Dict:
    payload = extract_text_from_any(path, filename)
    if not payload.get("text"):
        raise ValueError(f"No text extracted from {filename}")

    chunks = smart_chunk(payload["text"], target_tokens=300, overlap_tokens=40)
    n = upserter(user_id=user_id, source=filename, chunks=chunks, metadata=payload["meta"])
    return {"ok": True, "source": filename, "chunks_added": n, "meta": payload["meta"]}

from langchain.vectorstores import FAISS

def is_file_in_rag(vectorstore: FAISS, filename: str, k: int = 3) -> bool:
    """
    Check if a file was indexed in the RAG vectorstore.
    
    Args:
        vectorstore (FAISS): The loaded FAISS vectorstore instance.
        filename (str): The filename or partial path to search for in metadata.
        k (int): How many results to pull from the index for verification.
    
    Returns:
        bool: True if filename found in metadata, else False.
    """
    # Run a metadata search
    results = vectorstore.similarity_search(filename, k=k)
    
    for doc in results:
        meta = doc.metadata
        if "source" in meta and filename.lower() in meta["source"].lower():
            print(f"✅ File '{filename}' found in index. Metadata: {meta}")
            return True
    
    print(f"❌ File '{filename}' not found in index.")
    return False

